import pandas as pd
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import numpy as np

# =====================================================================
# CONFIGURATION: GROUP DEFINITIONS
# =====================================================================
# Define groups as a list of (NAME, DIRECTORY) tuples
# N_GROUPS will be automatically determined from the length of this list
# Add or remove groups as needed

BASE_DIR = Path(__file__).resolve().parents[1] / "rawdata"
EXPERIMENT_NAME = "Noise vs No Noise"

GROUPS_CONFIG = [
    ("No noise", "2026-01-07_12-43-42_no_noise"),
    ("N = 0.1, W = 1.0", "2026-01-07_12-43-09_noise_0_1"),
    ("N = 0.5, W = 1.0", "2026-01-07_13-37-40_noise_0_5"),
    ("N = 1.0, W = 1.0", "2026-01-07_13-38-29_noise_1_0"),
    ("Random", "2026-01-08_17-37-19_random"),
    ("N = 0.1, W = 0.6", "2026-01-08_19-45-52_weights_0_6"),
    # Add more groups here as needed
    # ("Group 5", "folder_5"),
    # ("Group 6", "folder_6"),
]

# Automatically determine number of groups
N_GROUPS = len(GROUPS_CONFIG)

# Base color palette: green, gold, cool steel, dark wine, brown, dark purple
BASE_COLORS = [
    "#0B3D2E",   # dark green
    "#D4AF37",   # gold
    "#88A0A8",   # cool steel
    "#721817",   # dark wine
    "#654236",   # brown
    "#331832",   # dark purple
]

print("BASE_DIR =", BASE_DIR)
print("Exists:", BASE_DIR.exists())
print("Contents:", list(BASE_DIR.iterdir()))


# =====================================================================
# Build active groups list and assign colors
# =====================================================================
GROUPS = []
PALETTE_COLORS = {}

for i in range(min(N_GROUPS, len(GROUPS_CONFIG))):
    name, dir_name = GROUPS_CONFIG[i]
    if name.strip():  # Skip empty group names
        GROUPS.append((name, BASE_DIR / dir_name))
        
        # Assign color from base palette or random if exhausted
        if i < len(BASE_COLORS):
            PALETTE_COLORS[name] = BASE_COLORS[i]
        else:
            # Generate random color for additional groups
            random_color = "#{:06x}".format(np.random.randint(0, 0xFFFFFF))
            PALETTE_COLORS[name] = random_color

print(f"\nActive groups: {[g[0] for g in GROUPS]}")
print(f"Colors assigned: {PALETTE_COLORS}")


# =====================================================================
# Helper function: load one condition (summary + runs)
# =====================================================================
def load_condition(condition_dir, condition_label):
    condition_dir = Path(condition_dir)

    # ---- load summary ----
    summary_file = next(condition_dir.glob("summary_*.csv"))
    df_summary = pd.read_csv(summary_file)
    df_summary["condition"] = condition_label

    # ---- load runs ----
    runs_dir = condition_dir / "runs"
    run_dfs = []

    for run_file in sorted(runs_dir.glob("run_*.csv")):
        df = pd.read_csv(run_file)

        # extract run id: run_0010.csv -> 10
        run_id = int(run_file.stem.split("_")[1])

        df["run_id"] = run_id
        df["condition"] = condition_label

        run_dfs.append(df)

    df_runs = pd.concat(run_dfs, ignore_index=True)

    return df_summary, df_runs


# =====================================================================
# Load all active conditions
# =====================================================================
all_summaries = []
all_runs = []

for group_name, group_dir in GROUPS:
    summary, runs = load_condition(group_dir, group_name)
    all_summaries.append(summary)
    all_runs.append(runs)

df_summary = pd.concat(all_summaries, ignore_index=True)
df_runs = pd.concat(all_runs, ignore_index=True)


# =====================================================================
# Sanity prints 
# =====================================================================
print("Summary shape:", df_summary.shape)
print("Runs shape:", df_runs.shape)
print(df_runs.head())


# =====================================================================
# Exploratory data analysis
# =====================================================================
import matplotlib.pyplot as plt
import seaborn as sns
from matplotlib.lines import Line2D

sns.set_theme(style="whitegrid", context="talk")

# Build palette for active groups only
palette = {name: PALETTE_COLORS[name] for name, _ in GROUPS}

plt.ion()


# Store figures for report generation
fig_paths = []
ks_test_results = []
FIGURES_DIR = Path(__file__).resolve().parent / "figures"
FIGURES_DIR.mkdir(exist_ok=True)

# --- energy over time
fig = plt.figure(figsize=(8, 5))

# Iterate through groups in defined order
for group_name, _ in GROUPS:
    for _, df_r in df_runs[df_runs["condition"] == group_name].groupby("run_id"):
        plt.plot(
            df_r["tick"],
            df_r["energy"],
            color=palette[group_name],
            alpha=0.4,
            linewidth=1
        )

plt.xlabel("Tick")
plt.ylabel("Energy")
plt.title("Energy dynamics per run")

legend_elements = [
    Line2D([0], [0], color=palette[name], lw=3, label=name)
    for name, _ in GROUPS
]

plt.legend(handles=legend_elements, frameon=False)
plt.tight_layout()

# Save figure for report
fig_path = FIGURES_DIR / "01_energy_dynamics.png"
fig.savefig(fig_path, dpi=150, bbox_inches="tight")
fig_paths.append(("Energy dynamics per run", fig_path))

plt.show()


# --- food over time
fig = plt.figure(figsize=(8, 5))

# Iterate through groups in defined order
for group_name, _ in GROUPS:
    for _, df_r in df_runs[df_runs["condition"] == group_name].groupby("run_id"):
        plt.plot(
            df_r["tick"],
            df_r["eats"],
            color=palette[group_name],
            alpha=0.4,
            linewidth=1
        )

plt.xlabel("Tick")
plt.ylabel("Food eaten")
plt.title("Food accumulation per run")

plt.legend(handles=legend_elements, frameon=False)
plt.tight_layout()

# Save figure for report
fig_path = FIGURES_DIR / "02_food_accumulation.png"
fig.savefig(fig_path, dpi=150, bbox_inches="tight")
fig_paths.append(("Food accumulation per run", fig_path))

plt.show()


# --- distance over time
fig = plt.figure(figsize=(8, 5))

# Iterate through groups in defined order
for group_name, _ in GROUPS:
    for _, df_r in df_runs[df_runs["condition"] == group_name].groupby("run_id"):
        plt.plot(
            df_r["tick"],
            df_r["distance"],
            color=palette[group_name],
            alpha=0.4,
            linewidth=1
        )

plt.xlabel("Tick")
plt.ylabel("Distance")
plt.title("Distance accumulation per run")

plt.legend(handles=legend_elements, frameon=False)
plt.tight_layout()

# Save figure for report
fig_path = FIGURES_DIR / "03_distance_accumulation.png"
fig.savefig(fig_path, dpi=150, bbox_inches="tight")
fig_paths.append(("Distance accumulation per run", fig_path))

plt.show()


# --- survival race plot
fig = plt.figure(figsize=(8, 5))

# Iterate through groups in defined order
for group_name, _ in GROUPS:
    df_c = df_summary[df_summary["condition"] == group_name]
    survival_times = df_c["lifetime_ticks"].values

    max_t = survival_times.max()
    ticks = range(max_t + 1)

    alive = [(survival_times >= t).sum() for t in ticks]

    plt.plot(
        ticks,
        alive,
        color=palette[group_name],
        linewidth=3,
        label=group_name
    )

plt.xlabel("Tick")
plt.ylabel("Bytes alive")
plt.title("Survival race of Byte simulations")

plt.legend(frameon=False)
plt.tight_layout()

# Save figure for report
fig_path = FIGURES_DIR / "04_survival_race.png"
fig.savefig(fig_path, dpi=150, bbox_inches="tight")
fig_paths.append(("Survival race", fig_path))

plt.show()


# --- KS test (pairwise comparisons for all groups)
from scipy.stats import ks_2samp

print("\n" + "="*60)
print("STATISTICAL TESTS (Kolmogorov-Smirnov)")
print("="*60)

group_names = [name for name, _ in GROUPS]
group_data = {
    name: df_summary.loc[df_summary["condition"] == name, "lifetime_ticks"].values
    for name, _ in GROUPS
}

for i, name1 in enumerate(group_names):
    for name2 in group_names[i+1:]:
        ks_stat, p_value = ks_2samp(group_data[name1], group_data[name2])
        print(f"\n{name1} vs {name2}:")
        print(f"  KS statistic: {ks_stat:.4f}")
        print(f"  p-value: {p_value:.4e}")
        ks_test_results.append((name1, name2, ks_stat, p_value))


# --- log-log CCDF
import numpy as np

fig = plt.figure(figsize=(8, 5))

# Iterate through groups in defined order
for group_name, _ in GROUPS:
    df_c = df_summary[df_summary["condition"] == group_name]
    lifetimes = np.sort(df_c["lifetime_ticks"].values)
    unique_t = np.unique(lifetimes)
    ccdf = [(lifetimes >= t).mean() for t in unique_t]

    plt.plot(
        unique_t,
        ccdf,
        marker="o",
        linestyle="none",
        markersize=4,
        alpha=0.7,
        color=palette[group_name],
        label=group_name
    )

plt.xscale("log")
plt.yscale("log")
plt.xlabel("Lifetime (ticks)")
plt.ylabel("P(T ≥ t)")
plt.title("Survival CCDF (log–log)")

plt.legend(frameon=False)
plt.tight_layout()

# Save figure for report
fig_path = FIGURES_DIR / "05_survival_ccdf_loglog.png"
fig.savefig(fig_path, dpi=150, bbox_inches="tight")
fig_paths.append(("Survival CCDF (log-log)", fig_path))

plt.show()


# =====================================================================
# Generate Report Document
# =====================================================================

doc = Document()

# Title
title = doc.add_paragraph()
title_run = title.add_run(f"Analysis Report: {EXPERIMENT_NAME}")
title_run.font.size = Pt(18)
title_run.bold = True
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Data folders section
doc.add_paragraph("Data Folders Used", style="Heading 2")
for group_name, group_dir in GROUPS:
    doc.add_paragraph(f"{group_name}: {group_dir}", style="List Bullet")

# KS test results section
doc.add_paragraph("Kolmogorov-Smirnov Test Results", style="Heading 2")
for name1, name2, ks_stat, p_value in ks_test_results:
    doc.add_paragraph(f"{name1} vs {name2}", style="Heading 3")
    table = doc.add_table(rows=3, cols=2)
    table.style = "Light Grid Accent 1"
    
    cells = table.rows[0].cells
    cells[0].text = "Metric"
    cells[1].text = "Value"
    
    cells = table.rows[1].cells
    cells[0].text = "KS statistic"
    cells[1].text = f"{ks_stat:.4f}"
    
    cells = table.rows[2].cells
    cells[0].text = "p-value"
    cells[1].text = f"{p_value:.4e}"

# Figures section
doc.add_paragraph("Figures", style="Heading 2")
for fig_title, fig_path in fig_paths:
    doc.add_paragraph(fig_title, style="Heading 3")
    doc.add_picture(str(fig_path), width=Inches(6))
    # Add some space after figure
    doc.add_paragraph()

# Save document
report_path = Path(__file__).resolve().parent / f"Report_{EXPERIMENT_NAME.replace(' ', '_')}.docx"
doc.save(report_path)
print(f"\nReport saved to: {report_path}")
print(f"Figures saved to: {FIGURES_DIR}")

input("Press Enter to close all figures...")
plt.ioff()
